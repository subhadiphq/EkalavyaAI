"""
EkalavyaAI — Premium PDF Notes Generator
Uses WeasyPrint to generate beautiful, exam-ready PDF notes.

Features:
- Handwriting-style font (Caveat from Google Fonts)
- Ruled notebook background lines
- EkalavyaAI watermark logo
- Color-coded importance markers
- Inline SVG diagrams
- PYQ Alert boxes
- Formula boxes with border
- Page numbers and chapter header
"""
import asyncio
import io
import logging
import os
import uuid
from typing import Optional

from config import settings

logger = logging.getLogger(__name__)

# ─── CSS Template ─────────────────────────────────────────────────────────────
NOTES_CSS = """
@import url('https://fonts.googleapis.com/css2?family=Caveat:wght@400;600;700&family=Inter:wght@400;500;600&display=swap');

:root {
  --blue: #1e40af;
  --amber: #d97706;
  --green: #065f46;
  --red: #dc2626;
  --gray: #6b7280;
  --light-gray: #f3f4f6;
  --border: #e5e7eb;
  --paper: #fefce8;
  --line-color: #dbeafe;
}

@page {
  size: A4;
  margin: 20mm 15mm 20mm 20mm;
  @top-center {
    content: element(header);
    font-family: 'Inter', sans-serif;
    font-size: 9pt;
    color: var(--gray);
  }
  @bottom-center {
    content: counter(page) " / " counter(pages);
    font-family: 'Inter', sans-serif;
    font-size: 9pt;
    color: var(--gray);
  }
}

body {
  font-family: 'Caveat', cursive;
  font-size: 14pt;
  line-height: 2;
  background-color: var(--paper);
  color: #1f2937;
  background-image: repeating-linear-gradient(
    transparent,
    transparent 31px,
    var(--line-color) 31px,
    var(--line-color) 32px
  );
  background-attachment: local;
}

/* Header */
.page-header {
  position: running(header);
  display: flex;
  justify-content: space-between;
  align-items: center;
  border-bottom: 1px solid var(--border);
  padding-bottom: 4px;
}

/* Title page */
.title-page {
  text-align: center;
  padding: 40px 0;
  page-break-after: always;
}

.title-page .logo {
  font-family: 'Inter', sans-serif;
  font-size: 28pt;
  font-weight: 700;
  color: var(--blue);
  letter-spacing: -1px;
}

.title-page .subtitle {
  font-size: 14pt;
  color: var(--amber);
  margin: 8px 0;
}

.title-page .chapter-name {
  font-size: 22pt;
  font-weight: 700;
  color: #1f2937;
  margin: 30px 0 10px;
  border-bottom: 3px solid var(--blue);
  padding-bottom: 8px;
}

.title-page .exam-badge {
  display: inline-block;
  background: var(--blue);
  color: white;
  font-family: 'Inter', sans-serif;
  font-size: 11pt;
  padding: 6px 16px;
  border-radius: 20px;
  margin: 10px 4px;
}

/* Section headings */
h1, h2, h3 {
  font-family: 'Inter', sans-serif;
  page-break-after: avoid;
}

h1 {
  font-size: 18pt;
  color: var(--blue);
  border-left: 4px solid var(--blue);
  padding-left: 12px;
  margin-top: 24px;
}

h2 {
  font-size: 15pt;
  color: #374151;
  margin-top: 16px;
}

h3 {
  font-size: 13pt;
  color: var(--amber);
  margin-top: 12px;
}

/* PYQ Alert Box */
.pyq-alert {
  background: #fef3c7;
  border-left: 4px solid var(--amber);
  padding: 10px 14px;
  margin: 12px 0;
  border-radius: 0 6px 6px 0;
  font-family: 'Inter', sans-serif;
  font-size: 11pt;
}

.pyq-alert::before {
  content: "⚡ PYQ ALERT: ";
  font-weight: 700;
  color: var(--amber);
}

/* Formula Box */
.formula-box {
  background: #eff6ff;
  border: 2px solid #bfdbfe;
  border-radius: 8px;
  padding: 12px 16px;
  margin: 12px 0;
  font-family: 'Caveat', cursive;
  font-size: 15pt;
  text-align: center;
}

.formula-box::before {
  content: "📐 Formula";
  display: block;
  font-family: 'Inter', sans-serif;
  font-size: 9pt;
  font-weight: 600;
  color: var(--blue);
  text-align: left;
  margin-bottom: 6px;
}

/* Memory Tip */
.memory-tip {
  background: #f0fdf4;
  border-left: 4px solid #16a34a;
  padding: 10px 14px;
  margin: 12px 0;
  border-radius: 0 6px 6px 0;
  font-size: 13pt;
}

.memory-tip::before {
  content: "🧠 Memory Tip: ";
  font-weight: 700;
  color: #16a34a;
  font-family: 'Inter', sans-serif;
  font-size: 11pt;
}

/* Exam Warning / Common Mistake */
.exam-warning {
  background: #fff1f2;
  border-left: 4px solid var(--red);
  padding: 10px 14px;
  margin: 12px 0;
  border-radius: 0 6px 6px 0;
  font-size: 13pt;
}

.exam-warning::before {
  content: "⚠️ Common Mistake: ";
  font-weight: 700;
  color: var(--red);
  font-family: 'Inter', sans-serif;
  font-size: 11pt;
}

/* Exam Tip */
.exam-tip {
  background: #f0f9ff;
  border: 1px dashed #7dd3fc;
  border-radius: 6px;
  padding: 10px 14px;
  margin: 12px 0;
  font-size: 13pt;
}

.exam-tip::before {
  content: "✨ Exam Tip: ";
  font-weight: 700;
  color: #0369a1;
  font-family: 'Inter', sans-serif;
  font-size: 11pt;
}

/* Importance badges */
.badge-high {
  display: inline-block;
  background: #dc2626;
  color: white;
  font-family: 'Inter', sans-serif;
  font-size: 8pt;
  padding: 2px 8px;
  border-radius: 10px;
  margin-left: 8px;
  vertical-align: middle;
}

.badge-medium {
  display: inline-block;
  background: var(--amber);
  color: white;
  font-family: 'Inter', sans-serif;
  font-size: 8pt;
  padding: 2px 8px;
  border-radius: 10px;
  margin-left: 8px;
  vertical-align: middle;
}

/* Summary box */
.summary-box {
  background: #f9fafb;
  border: 2px solid var(--blue);
  border-radius: 10px;
  padding: 16px 20px;
  margin: 20px 0;
}

.summary-box h3 {
  font-family: 'Inter', sans-serif;
  color: var(--blue);
  margin-top: 0;
}

/* Diagram wrapper */
.diagram-wrapper {
  text-align: center;
  margin: 16px 0;
  page-break-inside: avoid;
}

.diagram-wrapper .diagram-title {
  font-family: 'Inter', sans-serif;
  font-size: 10pt;
  color: var(--gray);
  margin-bottom: 8px;
  font-style: italic;
}

/* Watermark */
.watermark {
  position: fixed;
  bottom: 40mm;
  right: 20mm;
  opacity: 0.06;
  font-family: 'Inter', sans-serif;
  font-size: 60pt;
  font-weight: 900;
  color: var(--blue);
  transform: rotate(-30deg);
  pointer-events: none;
}

/* Table styling */
table {
  width: 100%;
  border-collapse: collapse;
  margin: 12px 0;
  font-family: 'Inter', sans-serif;
  font-size: 11pt;
}

th {
  background: var(--blue);
  color: white;
  padding: 8px 12px;
  text-align: left;
}

td {
  border: 1px solid var(--border);
  padding: 7px 12px;
}

tr:nth-child(even) td {
  background: #f8fafc;
}
"""


def _build_html(notes: dict, diagrams: list, student_name: str) -> str:
    """Build complete HTML for WeasyPrint rendering."""
    chapter_name = notes.get("chapter_name", "Chapter")
    exam = notes.get("exam", "CA")
    subject = notes.get("subject", "")
    language = notes.get("language", "English")
    sections = notes.get("sections", []) or notes.get("section", [])
    metadata = notes.get("metadata", {})
    personalization = notes.get("personalization", "")

    # Diagrams by trigger for easy lookup
    diagram_map = {d["trigger"].lower(): d for d in diagrams}

    sections_html = ""

    if personalization:
        sections_html += f"""
        <div class="exam-tip">{personalization}</div>
        """

    for section in sections:
        sec_type = section.get("type", "topic")
        title = section.get("title", "")
        content = section.get("content", "")
        importance = section.get("importance", "MEDIUM")
        pyq_alert = section.get("pyq_alert")
        formulas = section.get("formulas", [])
        memory_tips = section.get("memory_tips", [])
        examples = section.get("examples", [])
        mistakes = section.get("mistakes", [])
        tips = section.get("tips", [])
        points = section.get("points", [])

        badge = ""
        if importance == "HIGH":
            badge = '<span class="badge-high">HIGH PRIORITY</span>'
        elif importance == "MEDIUM":
            badge = '<span class="badge-medium">MEDIUM</span>'

        if sec_type == "opening_hook":
            sections_html += f"<h1>{title}</h1>\n<p>{content}</p>\n"

        elif sec_type == "topic":
            sections_html += f"<h1>{title} {badge}</h1>\n"
            if pyq_alert:
                sections_html += f'<div class="pyq-alert">{pyq_alert}</div>\n'
            sections_html += f"<p>{content}</p>\n"
            for formula in formulas:
                sections_html += f'<div class="formula-box">{formula}</div>\n'
            for tip in memory_tips:
                sections_html += f'<div class="memory-tip">{tip}</div>\n'
            for example in examples:
                ex_title = example.get("title", "Example")
                ex_content = example.get("content", "")
                sections_html += f"<h3>Example: {ex_title}</h3>\n<p>{ex_content}</p>\n"

            # Check if any diagram matches this topic
            for trigger, diagram in diagram_map.items():
                if trigger in title.lower() or trigger in content.lower():
                    sections_html += f"""
                    <div class="diagram-wrapper">
                      <div class="diagram-title">{diagram.get('title', 'Diagram')}</div>
                      {diagram.get('svg', '')}
                    </div>
                    """

        elif sec_type == "common_mistakes":
            sections_html += f"<h1>{title}</h1>\n"
            for mistake in mistakes:
                sections_html += f'<div class="exam-warning">{mistake}</div>\n'
            if content:
                sections_html += f"<p>{content}</p>\n"

        elif sec_type == "exam_tips":
            sections_html += f"<h1>{title}</h1>\n"
            for tip in tips:
                sections_html += f'<div class="exam-tip">{tip}</div>\n'

        elif sec_type == "summary":
            points_html = "".join(f"<li>{p}</li>" for p in points)
            sections_html += f"""
            <div class="summary-box">
              <h3>🔑 {title}</h3>
              <ul>{points_html}</ul>
            </div>
            """

    importance_score = metadata.get("importance_score", 0.7)
    stars = "★" * round(importance_score * 5) + "☆" * (5 - round(importance_score * 5))

    html = f"""
<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <style>{NOTES_CSS}</style>
</head>
<body>
  <div class="watermark">EkalavyaAI</div>

  <div class="page-header">
    <span><strong>EkalavyaAI</strong> — {exam}</span>
    <span>{chapter_name}</span>
    <span>{student_name}</span>
  </div>

  <!-- Title Page -->
  <div class="title-page">
    <div class="logo">EkalavyaAI</div>
    <div class="subtitle">Learn Like a Topper</div>
    <div class="chapter-name">{chapter_name}</div>
    <div>
      <span class="exam-badge">{exam}</span>
      <span class="exam-badge">{subject}</span>
      <span class="exam-badge">{language}</span>
    </div>
    <p style="font-family: Inter; font-size: 12pt; color: #6b7280; margin-top: 20px;">
      Prepared for: <strong>{student_name}</strong><br>
      Importance: {stars} ({round(importance_score * 100)}%)<br>
      Expected exam marks: {metadata.get('estimated_exam_marks', '?')} marks
    </p>
  </div>

  <!-- Main Notes Content -->
  <div class="notes-content">
    {sections_html}
  </div>

</body>
</html>
"""
    return html


class PDFGenerator:
    """Premium PDF notes generator using WeasyPrint."""

    async def generate(
        self,
        notes: dict,
        diagrams: list,
        student_name: str,
        chapter_name: str,
        exam: str,
        language: str,
    ) -> tuple[str, str]:
        """
        Generate PDF and DOCX from structured notes.
        Returns (pdf_url, docx_url).
        """
        doc_id = str(uuid.uuid4())[:8]
        safe_chapter = chapter_name.replace(" ", "_").replace("/", "-")[:50]
        base_filename = f"{exam}_{safe_chapter}_{doc_id}"

        # Generate PDF in executor (WeasyPrint is sync)
        pdf_url = await asyncio.get_event_loop().run_in_executor(
            None, self._generate_pdf_sync, notes, diagrams, student_name, base_filename
        )

        # Generate DOCX
        docx_url = await self._generate_docx(notes, student_name, base_filename)

        return pdf_url, docx_url

    def _generate_pdf_sync(
        self,
        notes: dict,
        diagrams: list,
        student_name: str,
        base_filename: str,
    ) -> str:
        """Synchronous PDF generation — runs in thread executor."""
        try:
            from weasyprint import HTML, CSS

            html_content = _build_html(notes, diagrams, student_name)
            pdf_bytes = HTML(
                string=html_content,
                base_url="https://fonts.googleapis.com",
            ).write_pdf(
                presentational_hints=True,
            )

            # Upload to S3/R2
            pdf_url = self._upload_to_storage(
                file_bytes=pdf_bytes,
                filename=f"notes/{base_filename}.pdf",
                content_type="application/pdf",
            )
            return pdf_url
        except Exception as e:
            logger.error(f"[PDFGenerator] PDF generation error: {e}")
            return ""

    async def _generate_docx(
        self, notes: dict, student_name: str, base_filename: str
    ) -> str:
        """Generate DOCX using python-docx."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Title
            title = doc.add_heading(notes.get("chapter_name", "Notes"), 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph(f"Exam: {notes.get('exam')} | For: {student_name}")

            for section in notes.get("sections", []):
                sec_type = section.get("type", "topic")
                if sec_type in ("topic", "opening_hook"):
                    doc.add_heading(section.get("title", ""), level=1)
                    doc.add_paragraph(section.get("content", ""))
                    for formula in section.get("formulas", []):
                        p = doc.add_paragraph(f"📐 {formula}")
                        p.runs[0].bold = True
                elif sec_type == "summary":
                    doc.add_heading("Quick Revision Summary", level=1)
                    for point in section.get("points", []):
                        doc.add_paragraph(point, style="List Bullet")

            # Save to bytes
            docx_io = io.BytesIO()
            doc.save(docx_io)
            docx_bytes = docx_io.getvalue()

            return self._upload_to_storage(
                file_bytes=docx_bytes,
                filename=f"notes/{base_filename}.docx",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except Exception as e:
            logger.error(f"[PDFGenerator] DOCX generation error: {e}")
            return ""

    def _upload_to_storage(
        self, file_bytes: bytes, filename: str, content_type: str
    ) -> str:
        """Upload file to S3 or Cloudflare R2."""
        try:
            import boto3

            if settings.USE_R2:
                s3 = boto3.client(
                    "s3",
                    endpoint_url=settings.CLOUDFLARE_R2_ENDPOINT,
                    aws_access_key_id=settings.CLOUDFLARE_R2_ACCESS_KEY,
                    aws_secret_access_key=settings.CLOUDFLARE_R2_SECRET_KEY,
                )
                bucket = "ekalavya-notes-r2"
                base_url = settings.CLOUDFLARE_R2_ENDPOINT
            else:
                s3 = boto3.client(
                    "s3",
                    aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
                    aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY,
                    region_name=settings.AWS_REGION,
                )
                bucket = settings.AWS_S3_BUCKET
                base_url = f"https://{bucket}.s3.{settings.AWS_REGION}.amazonaws.com"

            s3.put_object(
                Bucket=bucket,
                Key=filename,
                Body=file_bytes,
                ContentType=content_type,
                ACL="private",  # Signed URLs for access
            )

            # Generate pre-signed URL (24 hours)
            url = s3.generate_presigned_url(
                "get_object",
                Params={"Bucket": bucket, "Key": filename},
                ExpiresIn=86400,
            )
            return url
        except Exception as e:
            logger.error(f"[PDFGenerator] Storage upload error: {e}")
            return f"/local/{filename}"  # Fallback for development
